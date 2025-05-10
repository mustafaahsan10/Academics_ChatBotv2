import os
import glob
import pandas as pd
import re
from pathlib import Path
from dotenv import load_dotenv
import argparse
from PyPDF2 import PdfReader
from docx import Document
import csv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct
import nltk
import datetime

# Load environment variables
load_dotenv()

# Download NLTK data for tokenization
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Configure parser
parser = argparse.ArgumentParser(description="Ingest university data into Qdrant vector database")
parser.add_argument("--data_dir", default="data/raw", help="Directory containing data files")
parser.add_argument("--chunk_size", type=int, default=1000, help="Size of document chunks")
parser.add_argument("--chunk_overlap", type=int, default=200, help="Overlap between document chunks")
parser.add_argument("--recreate", action="store_true", help="Recreate the collection if it exists")
args = parser.parse_args()

# Regex patterns for identifying document structure
HEADING_PATTERN = r'^(?:[A-Z0-9][\.\)]\s+)?[A-Z][A-Za-z0-9\s:]{1,60}$'
COURSE_CODE_PATTERN = r'\b([A-Z]{2,4})\s?(\d{3}[A-Z]?)\b'

def extract_headings_from_docx(doc):
    """Extract headings from docx document with style information"""
    headings = []
    current_heading = None
    current_subheading = None
    
    for para in doc.paragraphs:
        # Check if paragraph is a heading
        if para.style.name.startswith('Heading 1') and para.text.strip():
            current_heading = para.text.strip()
            current_subheading = None
            headings.append({
                "level": 1,
                "text": current_heading,
                "parent": None,
                "children": []
            })
        elif para.style.name.startswith('Heading 2') and para.text.strip():
            current_subheading = para.text.strip()
            if current_heading:
                # Find parent heading and add this as child
                for h in headings:
                    if h["level"] == 1 and h["text"] == current_heading:
                        h["children"].append(current_subheading)
            headings.append({
                "level": 2,
                "text": current_subheading,
                "parent": current_heading,
                "children": []
            })
        elif para.style.name.startswith('Heading') and para.text.strip():
            # Other heading levels
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 3
            parent = current_subheading if level > 2 and current_subheading else current_heading
            headings.append({
                "level": level,
                "text": para.text.strip(),
                "parent": parent,
                "children": []
            })
    
    return headings

def extract_headings_from_text(text_lines):
    """Extract headings from plain text using heuristics"""
    headings = []
    current_heading = None
    current_subheading = None
    
    for i, line in enumerate(text_lines):
        line = line.strip()
        if not line:
            continue
        
        # Check if line might be a heading
        is_heading = False
        
        # Heading patterns:
        # 1. All caps with reasonable length
        if line.isupper() and 3 < len(line) < 60:
            is_heading = True
            level = 1
        # 2. Matches heading pattern (starts with cap, reasonable length)
        elif re.match(HEADING_PATTERN, line):
            is_heading = True
            level = 2 if current_heading else 1
        # 3. Numbered or lettered section (e.g., "1. Introduction")
        elif re.match(r'^[0-9]+\.\s+[A-Z]', line) or re.match(r'^[A-Z]\.\s+[A-Z]', line):
            is_heading = True
            level = 2 if current_heading else 1
        # 4. Contains a course code at the beginning
        elif re.match(r'^' + COURSE_CODE_PATTERN, line):
            is_heading = True
            level = 2
        
        if is_heading:
            if level == 1:
                current_heading = line
                current_subheading = None
                headings.append({
                    "level": 1,
                    "text": current_heading,
                    "parent": None,
                    "children": []
                })
            else:
                current_subheading = line
                if current_heading:
                    # Find parent heading and add this as child
                    for h in headings:
                        if h["level"] == 1 and h["text"] == current_heading:
                            h["children"].append(current_subheading)
                headings.append({
                    "level": 2,
                    "text": current_subheading,
                    "parent": current_heading,
                    "children": []
                })
    
    return headings

def find_relevant_heading(chunk_text, all_headings, all_text):
    """Find the most relevant heading for a chunk of text"""
    # Find position of chunk in the full text
    chunk_position = all_text.find(chunk_text)
    if chunk_position == -1:
        return None, None
    
    # Find what text precedes this chunk
    preceding_text = all_text[:chunk_position]
    
    # Find the closest heading before this position
    current_heading = None
    current_subheading = None
    
    for heading in all_headings:
        heading_pos = preceding_text.find(heading["text"])
        if heading_pos != -1:
            if heading["level"] == 1:
                current_heading = heading["text"]
                current_subheading = None
            else:
                current_subheading = heading["text"]
    
    return current_heading, current_subheading

def create_metadata_for_chunk(chunk, doc_title, source_file, doc_type, all_headings, all_text):
    """Create metadata for a specific chunk with proper heading attribution"""
    # Find relevant heading for this chunk
    heading, subheading = find_relevant_heading(chunk, all_headings, all_text)
    
    # Extract relevant keywords from the chunk
    keywords = []
    if "prerequisite" in chunk.lower() or "prereq" in chunk.lower():
        keywords.append("prerequisites")
    if "syllabus" in chunk.lower():
        keywords.append("syllabus")
    if "exam" in chunk.lower():
        keywords.append("exam")
    if "assignment" in chunk.lower():
        keywords.append("assignment")
    if "grade" in chunk.lower():
        keywords.append("grade")
    if "credit" in chunk.lower():
        keywords.append("credit")
    if "schedule" in chunk.lower() or "timetable" in chunk.lower():
        keywords.append("schedule")
    
    # Extract course codes
    course_codes = re.findall(COURSE_CODE_PATTERN, chunk)
    entities = [f"{code[0]} {code[1]}" for code in course_codes]
    
    # Add common academic terms if present
    academic_terms = ["data structures", "algorithms", "programming", "database", 
                      "software engineering", "artificial intelligence", "networking"]
    for term in academic_terms:
        if term.lower() in chunk.lower() and term not in entities:
            entities.append(term)
    
    # Create metadata object
    metadata = {
        "doc_title": doc_title,
        "section": heading,
        "heading": subheading or heading,
        "entities": entities[:5],  # Limit to 5 entities
        "type": doc_type,
        "keywords": keywords,
        "source": str(source_file) if source_file else None
    }
    
    return metadata

def read_pdf(file_path):
    """Extract text and structured headings from PDF file"""
    try:
        pdf_reader = PdfReader(file_path)
        
        # Extract document title
        doc_title = Path(file_path).stem
        if pdf_reader.metadata and hasattr(pdf_reader.metadata, 'title') and pdf_reader.metadata.title:
            doc_title = pdf_reader.metadata.title.strip()
        
        # Extract text from all pages
        full_text = ""
        for page in pdf_reader.pages:
            full_text += page.extract_text() + "\n"
        
        # Extract headings from text
        text_lines = full_text.split('\n')
        headings = extract_headings_from_text(text_lines)
        
        return full_text, doc_title, headings, "course_document"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return "", "", [], ""

def read_docx(file_path):
    """Extract text and structured headings from DOCX file"""
    try:
        doc = Document(file_path)
        
        # Extract document title
        doc_title = Path(file_path).stem
        if doc.core_properties.title:
            doc_title = doc.core_properties.title.strip()
        
        # Extract text with paragraph markers
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        # Extract headings using style information
        headings = extract_headings_from_docx(doc)
        
        # If no styled headings found, try text-based extraction
        if not headings:
            text_lines = full_text.split('\n')
            headings = extract_headings_from_text(text_lines)
        
        return full_text, doc_title, headings, "course_document"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return "", "", [], ""

def read_txt(file_path):
    """Read text file and extract headings"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()
        
        # Extract document title from filename
        doc_title = Path(file_path).stem
        
        # Extract headings
        text_lines = full_text.split('\n')
        headings = extract_headings_from_text(text_lines)
        
        return full_text, doc_title, headings, "text_document"
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return "", "", [], ""

def read_excel_or_csv(file_path):
    """Extract structured data from Excel or CSV file"""
    try:
        # Determine file type
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.csv':
            # Read CSV
            df = pd.read_csv(file_path)
        else:
            # Read Excel
            df = pd.read_excel(file_path)
        
        # Extract document title from filename
        doc_title = Path(file_path).stem
        
        # Convert to structured text format
        full_text = f"Data from: {Path(file_path).name}\n\n"
        
        # Use column names as headings
        headings = []
        for col in df.columns:
            headings.append({
                "level": 1,
                "text": col,
                "parent": None,
                "children": []
            })
        
        # Add column information
        full_text += f"Table contains {len(df)} rows and {len(df.columns)} columns.\n"
        full_text += "Columns: " + ", ".join(df.columns) + "\n\n"
        
        # Add sample rows
        if len(df) > 0:
            max_rows = min(10, len(df))
            full_text += f"Sample data (showing {max_rows} of {len(df)} rows):\n"
            for i, row in df.head(max_rows).iterrows():
                full_text += f"Row {i+1}: " + " | ".join([f"{col}: {str(val).strip()}" for col, val in row.items()]) + "\n"
        
        return full_text, doc_title, headings, "tabular_data"
    except Exception as e:
        print(f"Error reading Excel/CSV {file_path}: {e}")
        return "", "", [], ""

def process_directory():
    """Process all files in the data directory with improved heading extraction"""
    documents = []
    doc_titles = []
    all_headings = []
    doc_types = []
    file_paths = []
    
    # Get all files
    data_dir = Path(args.data_dir)
    pdf_files = list(data_dir.glob("**/*.pdf"))
    docx_files = list(data_dir.glob("**/*.docx"))
    txt_files = list(data_dir.glob("**/*.txt"))
    excel_files = list(data_dir.glob("**/*.xlsx")) + list(data_dir.glob("**/*.xls"))
    csv_files = list(data_dir.glob("**/*.csv"))
    
    # Process PDFs
    for file_path in pdf_files:
        print(f"Processing PDF: {file_path}")
        text, title, headings, doc_type = read_pdf(file_path)
        if text:
            documents.append(text)
            doc_titles.append(title)
            all_headings.append(headings)
            doc_types.append(doc_type)
            file_paths.append(file_path)
    
    # Process DOCXs
    for file_path in docx_files:
        print(f"Processing DOCX: {file_path}")
        text, title, headings, doc_type = read_docx(file_path)
        if text:
            documents.append(text)
            doc_titles.append(title)
            all_headings.append(headings)
            doc_types.append(doc_type)
            file_paths.append(file_path)
    
    # Process TXTs
    for file_path in txt_files:
        print(f"Processing TXT: {file_path}")
        text, title, headings, doc_type = read_txt(file_path)
        if text:
            documents.append(text)
            doc_titles.append(title)
            all_headings.append(headings)
            doc_types.append(doc_type)
            file_paths.append(file_path)
            
    # Process Excel files
    for file_path in excel_files:
        print(f"Processing Excel: {file_path}")
        text, title, headings, doc_type = read_excel_or_csv(file_path)
        if text:
            documents.append(text)
            doc_titles.append(title)
            all_headings.append(headings)
            doc_types.append(doc_type)
            file_paths.append(file_path)
    
    # Process CSV files
    for file_path in csv_files:
        print(f"Processing CSV: {file_path}")
        text, title, headings, doc_type = read_excel_or_csv(file_path)
        if text:
            documents.append(text)
            doc_titles.append(title)
            all_headings.append(headings)
            doc_types.append(doc_type)
            file_paths.append(file_path)
    
    return documents, doc_titles, all_headings, doc_types, file_paths

def chunk_documents(documents, doc_titles, all_headings, doc_types, file_paths):
    """Split documents into chunks with heading-aware metadata"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap
    )
    
    all_chunks = []
    all_metadatas = []
    
    for i, (doc, title, headings, doc_type, file_path) in enumerate(zip(documents, doc_titles, all_headings, doc_types, file_paths)):
        # Split the document into chunks
        chunks = text_splitter.split_text(doc)
        
        # Create metadata for each chunk based on its content and position
        chunk_metadatas = []
        for chunk in chunks:
            metadata = create_metadata_for_chunk(chunk, title, file_path, doc_type, headings, doc)
            chunk_metadatas.append(metadata)
        
        all_chunks.extend(chunks)
        all_metadatas.extend(chunk_metadatas)
    
    return all_chunks, all_metadatas

def setup_qdrant():
    """Set up Qdrant collection with recreate option"""
    client = QdrantClient(
        url=os.getenv("QDRANT_URL"),
        api_key=os.getenv("QDRANT_API_KEY")
    )
    
    # Check if collection exists
    collections = client.get_collections().collections
    collection_names = [collection.name for collection in collections]
    
    # Delete collection if it exists and recreate flag is set
    if "university_data" in collection_names and args.recreate:
        print("Recreating collection 'university_data'...")
        client.delete_collection(collection_name="university_data")
        collection_exists = False
    else:
        collection_exists = "university_data" in collection_names
    
    # Create collection if it doesn't exist
    if not collection_exists:
        client.create_collection(
            collection_name="university_data",
            vectors_config=VectorParams(
                size=1536,  # OpenAI embedding dimension
                distance=Distance.COSINE
            )
        )
        print("Created new 'university_data' collection")
    else:
        print("Collection 'university_data' already exists")
    
    return client

def generate_embeddings_and_upload(chunks, metadatas):
    """Generate embeddings and upload to Qdrant with heading-aware metadata"""
    # Initialize OpenAI embeddings
    embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY"))
    
    # Initialize Qdrant client
    client = setup_qdrant()
    
    print(f"Generating embeddings for {len(chunks)} chunks...")
    
    # Process in batches to avoid overloading the API
    batch_size = 10
    total_points = 0
    
    for i in range(0, len(chunks), batch_size):
        end_idx = min(i + batch_size, len(chunks))
        batch_chunks = chunks[i:end_idx]
        batch_metadatas = metadatas[i:end_idx]
        
        try:
            # Generate embeddings
            batch_embeddings = embeddings.embed_documents(batch_chunks)
            
            # Prepare points
            points = []
            for j, (chunk, metadata, embedding_vector) in enumerate(zip(batch_chunks, batch_metadatas, batch_embeddings)):
                # Create point with embedding, payload with metadata, and ID
                points.append(
                    PointStruct(
                        id=total_points + j,
                        vector=embedding_vector,
                        payload={
                            "text": chunk,
                            "metadata": metadata
                        }
                    )
                )
            
            # Upload to Qdrant
            client.upsert(
                collection_name="university_data",
                points=points
            )
            
            total_points += len(batch_chunks)
            print(f"Processed {total_points}/{len(chunks)} chunks")
            
        except Exception as e:
            print(f"Error processing batch {i}-{end_idx}: {e}")
    
    print(f"Successfully ingested {total_points} chunks into Qdrant")

def main():
    """Main ingestion process with improved heading extraction"""
    print("Starting data ingestion process...")
    
    # Process files
    print("Reading documents from directory...")
    documents, doc_titles, all_headings, doc_types, file_paths = process_directory()
    print(f"Found {len(documents)} documents")
    
    if not documents:
        print("No documents found. Please check the data directory.")
        return
    
    # Chunk documents
    print("Chunking documents...")
    chunks, metadatas = chunk_documents(documents, doc_titles, all_headings, doc_types, file_paths)
    print(f"Created {len(chunks)} chunks")
    
    # Upload to Qdrant
    print("Uploading to Qdrant...")
    generate_embeddings_and_upload(chunks, metadatas)
    
    print("Data ingestion complete!")

if __name__ == "__main__":
    main() 